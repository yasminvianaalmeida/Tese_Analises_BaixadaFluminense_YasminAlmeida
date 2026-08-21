# -*- coding: utf-8 -*-
# Gera os APENDICES (vocabulario, livro de codigo, rubrica) A PARTIR da planilha de parametros.
# Garante que texto e parametros NUNCA divergem. Carimba versao (SHA-256 + data).
import sys, hashlib, datetime, openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ARQ_PARAM = sys.argv[1] if len(sys.argv)>1 else "Parametros_LDO_LOA.xlsx"
ARQ_SAIDA = sys.argv[2] if len(sys.argv)>2 else "APENDICES_geradas_dos_parametros.docx"

def linhas(ws):
    out=[]
    for r in range(3, ws.max_row+1):
        vals=[ws.cell(r,c).value for c in range(1, ws.max_column+1)]
        if any(v not in (None,"") for v in vals): out.append(vals)
    return out

wb=openpyxl.load_workbook(ARQ_PARAM)
# SHA-256 do arquivo
sha=hashlib.sha256(open(ARQ_PARAM,"rb").read()).hexdigest()
data=datetime.date.today().isoformat()

termos_busca={}
for chave,padrao,*_ in linhas(wb["Termos_Busca"]):
    if chave and padrao: termos_busca.setdefault(str(chave).strip(),[]).append(str(padrao).strip())
tema={str(a).strip():str(b).strip() for a,b,*_ in linhas(wb["Tema_Padrao"]) if a}
forca={str(a).strip():int(b) for a,b,*_ in linhas(wb["Forca_Tecnica"]) if a and b not in (None,"")}
proj_termo={}
for projeto,eixo,ader,termos,interp,*_ in linhas(wb["Projetos"]):
    for t in str(termos or "").split(";"):
        if t.strip(): proj_termo[t.strip().lower()]=str(projeto).strip()
voc={}
for lista,valor,*_ in linhas(wb["Vocabulario_Apoio"]):
    if lista and valor: voc.setdefault(str(lista).strip(),[]).append(str(valor).strip())
grau=linhas(wb["Grau_Rubrica"])
ger={str(a).strip():b for a,b,*_ in linhas(wb["Parametros_Gerais"])}

doc=Document()
st=doc.styles["Normal"]; st.font.name="Times New Roman"; st.font.size=Pt(12)
def H(t,l=1):
    h=doc.add_heading(t,level=l)
    for r in h.runs: r.font.color.rgb=RGBColor(0x1F,0x38,0x64)
def P(t,it=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; r=p.add_run(t); r.italic=it; return p

t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("APÊNDICES – Vocabulário, livro de código e parâmetros"); r.bold=True; r.font.size=Pt(15)
P(f"Documento gerado automaticamente a partir da planilha de parâmetros, em {data}. "
  f"Arquivo-fonte: {ARQ_PARAM.split('/')[-1]}. Resumo criptográfico (SHA-256): {sha}.", it=True)
P("Em caso de divergência entre versões textuais e arquivos computacionais, prevalecem os "
  "parâmetros congelados nesta planilha, identificada pela data, versão e resumo criptográfico acima.", it=True)

# APENDICE A - vocabulario
H("APÊNDICE A – Vocabulário controlado")
tab=doc.add_table(rows=1, cols=5); tab.style="Light Grid Accent 1"
for i,h in enumerate(["Termo-chave","Padrões (regex)","Força","Tema","Projeto"]): tab.rows[0].cells[i].text=h
for termo in sorted(termos_busca):
    c=tab.add_row().cells
    c[0].text=termo; c[1].text=" ; ".join(termos_busca[termo])
    c[2].text=str(forca.get(termo,"")); c[3].text=tema.get(termo,""); c[4].text=proj_termo.get(termo.lower(),"")
    for cell in c:
        for p in cell.paragraphs:
            for rr in p.runs: rr.font.size=Pt(8)

# APENDICE B - livro de codigo
H("APÊNDICE B – Livro de código")
P("Fórmula do peso analítico: P = F × A, em que F é a força técnica do termo (1 a 3) e A é a "
  "ancoragem territorial do contexto (0 a 2). O peso varia de 0 a 6; são aproveitáveis as ocorrências com peso maior que zero.")
tiers={3:[],2:[],1:[],0:[]}
for k,v in forca.items(): tiers.setdefault(v,[]).append(k)
H("B.1 Força técnica por termo (congelada)",2)
for val in [3,2,1,0]:
    if tiers.get(val): P(f"Força {val}: " + ", ".join(sorted(tiers[val])) + ".")
H("B.2 Listas de apoio à classificação",2)
rot={"PALAVRAS_TERRITORIAIS":"Palavras territoriais","PALAVRAS_ADMINISTRATIVAS":"Palavras administrativas genéricas",
     "EXPRESSOES_FALSO_POSITIVO":"Expressões de falso positivo","EXPRESSOES_CADASTRO_FISCAL_FRACO":"Expressões de cadastro fiscal fraco"}
for k,nome in rot.items():
    if voc.get(k): P(f"{nome}: " + ", ".join(voc[k]) + ".")
H("B.3 Rubrica do Grau (congelada)",2)
tb=doc.add_table(rows=1,cols=6); tb.style="Light Grid Accent 1"
hdr=["Nível","Soma peso ≥","Aproveit. ≥","Anos ≥","Temas ≥","Evid. forte"]
for i,h in enumerate(hdr): tb.rows[0].cells[i].text=h
for row in grau:
    c=tb.add_row().cells
    for i in range(6): c[i].text=str(row[i]) if i<len(row) and row[i] is not None else ""

# APENDICE D - parametros gerais
H("APÊNDICE D – Parâmetros gerais de execução")
tb=doc.add_table(rows=1,cols=2); tb.style="Light Grid Accent 1"
tb.rows[0].cells[0].text="Parâmetro"; tb.rows[0].cells[1].text="Valor"
for k,v in ger.items():
    c=tb.add_row().cells; c[0].text=str(k); c[1].text=str(v)

doc.save(ARQ_SAIDA)
print("Apendices gerados:", ARQ_SAIDA)
print("SHA-256 dos parametros:", sha)
print("Forças congeladas:", {k:forca[k] for k in sorted(forca)})
